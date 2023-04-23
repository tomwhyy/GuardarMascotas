import sqlite3
import os
import time
import sys
import docx

from colorama import Fore, init
from os import system
from dotenv import load_dotenv
from colored import fg


class PetSaver:
    def __init__(self):
        super().__init__()
        
        load_dotenv('.env')
        self.DB_ROUTE = os.getenv('DB_NAME')
        init()
        self.check_db = True
        self.banner_cat = r"""
         ,_     _           
          |\\_,-~/            
         / _  _ |    ,--.       
        (  @  @ )   / ,-'       
         \  _T_/-._( (      
         /         `. \     
        |         _  \ |        
         \ \ ,  /      |    
          || |-_\__   /     
         ((_/`(____,-'      
        """
        self.banner_cat_1 = r""" ,_     _'"""
        self.banner_cat_2 = r"""  |\\_,-~/"""
        self.banner_cat_3 = r""" / _  _ |    ,--."""
        self.banner_cat_4 = r"""(  @  @ )   / ,-'"""
        self.banner_cat_5 = r""" \  _T_/-._( ("""
        self.banner_cat_6 = r""" /         `. \""""
        self.banner_cat_7 = r"""|         _  \ |"""
        self.banner_cat_8 = r""" \ \ ,  /      |"""
        self.banner_cat_9 = r"""  || |-_\__   /"""
        self.banner_cat_10 = r""" ((_/`(____,-'"""
        
        self.app_version = os.getenv('VERSION')
        self.black = Fore.BLACK
        self.error_color = Fore.RED
        self.green = Fore.GREEN
        self.yellow = Fore.YELLOW
        self.blue = Fore.BLUE
        self.magenta = Fore.MAGENTA
        self.cyan = Fore.CYAN
        self.white = Fore.WHITE
        self.reset = Fore.RESET

        blue = fg(6)
        reset = fg(7)
        self.red = fg(1)
        green = fg(2)
        self.purple = fg(5)

        self.lblack = Fore.LIGHTBLACK_EX
        self.lred = Fore.LIGHTRED_EX
        self.lgreen = Fore.LIGHTGREEN_EX
        self.lyellow = Fore.LIGHTYELLOW_EX
        self.lblue = Fore.LIGHTBLUE_EX
        self.lmagenta = Fore.LIGHTMAGENTA_EX
        self.lcyan = Fore.LIGHTCYAN_EX
        self.lwhite = Fore.LIGHTWHITE_EX
    def start(self):
        """
        Esta funcion verifica si la base de datos ya esta creada, de lo contrario
        hará una nueva.
        """
        if os.path.isfile(self.DB_ROUTE):
            pass
        else:
            print(f'{self.white}[{self.cyan}DB{self.white}] No database found, creating one for you :)')
            conexion = sqlite3.connect(self.DB_ROUTE)

            cursor = conexion.cursor()

            cursor.execute("CREATE TABLE IF NOT EXISTS Pets " \
                "(id INTEGER AUTO_INCREMENT PRIMARY KEY, nombre VARCHAR(100), edad INTEGER, especie VARCHAR(100), Genero VARCHAR(32), nombre_dueno VARCHAR(128))")

            conexion.commit()

            conexion.close()
            time.sleep(0.8)
            system('cls' if os.name == 'nt' else 'clear')
    def add_pet(self):
        """
        Menú de agregar mascotas en la BD.
        """

        system('cls' if os.name == 'nt' else 'clear')
        system('title petSaver ~ Add Pet menu' if os.name == 'nt' else '')

        time.sleep(0.2)
        self.banner_addPet = f"""

                {self.white}[{self.cyan}Add Pet Menu{self.white}]                                      {self.lcyan}{self.banner_cat_1}
                                                                  {self.lcyan}{self.banner_cat_2}
        {self.lyellow}Bienvenido, sigue las instrucciones                        {self.lcyan}{self.banner_cat_3}
        {self.lyellow}para agregar una mascota a la base de datos.                {self.lcyan}{self.banner_cat_4}
                                                                    {self.lcyan}{self.banner_cat_5}
                                                                    {self.lcyan}{self.banner_cat_6}
                                                                    {self.lcyan}{self.banner_cat_7}
                                                                    {self.lcyan}{self.banner_cat_8}
                                                                    {self.lcyan}{self.banner_cat_9}
                    [x] Atrás                                       {self.lcyan}{self.banner_cat_10}
        """
        print(self.banner_addPet)
        nombre_ = input(f'{self.white}[{self.lcyan}Nombre{self.white}] {self.reset}')
        if len(nombre_) < 3:
            if nombre_ == 'x':
                self.main()
            print(f'{self.error_color}[ERROR] Nombre inválido, mínimo 3 caracteres')
            time.sleep(0.5)
            self.add_pet()
        else:
            edad_ = input(f'\n{self.white}[{self.lcyan}Edad (Ej. 3 Meses){self.white}] {self.reset}')
            if edad_ == '':
                print(f'{self.error_color}[ERROR] Edad inválida, introduce un valor.')
                time.sleep(0.5)
                self.add_pet()
            else:
                especie_ = input(f'\n{self.white}[{self.lcyan}Especie{self.white}] {self.reset}')
                if especie_ == '':
                    print(f'{self.error_color}[ERROR] Especie inválida, introduce un valor.')
                    time.sleep(0.5)
                    self.add_pet()
                else:
                    genero_ = input(f'\n{self.white}[{self.lcyan}Género{self.white}] {self.reset}')
                    if genero_ == '':
                        print(f'{self.error_color}[ERROR] Género inválido, introduce un valor.')
                        time.sleep(0.5)
                        self.add_pet()
                    else:
                        n_dueno_ = input(f'\n{self.white}[{self.lcyan}Nombre del dueño{self.white}] {self.reset}')
                        time.sleep(0.5)
                        if n_dueno_ == '':
                            print(f'\n{self.error_color}[ERROR] Nombre inválido, introduce un valor.')
                            time.sleep(0.5) 
                            self.add_pet()
                        else:
                            self.save_pet(nombre_, edad_, especie_, genero_, n_dueno_)
                            time.sleep(3)
    def save_pet(self, nombre: str, edad: str, especie: str, genero: str, n_dueno: str):
        """
        Guardar una mascota en la BD.

        Todos los argumentos son necesarios.
        """
        if os.path.isfile(self.DB_ROUTE):

            conexion = sqlite3.connect(self.DB_ROUTE)
            cursor = conexion.cursor()
            cursor.execute(f'SELECT * FROM Pets WHERE nombre="{nombre}"')
            get_nombreDb = cursor.fetchone()

            if get_nombreDb:
                print(f'\n{self.white}[{self.cyan}!{self.white}] {self.lcyan}{nombre}{self.cyan} ya se encuentra registrado en la BD.{self.reset}')
                opt = input(f'{self.white}[{self.magenta}>{self.white}] {self.reset}').lower()
                if opt == 'x':
                    self.add_pet()
            else:
                try:
                    conexion = sqlite3.connect(self.DB_ROUTE)
                    cursor = conexion.cursor()

                    cursor.execute(f"INSERT INTO Pets VALUES " \
                        f"(Null, '{nombre}', '{edad}', '{especie}', '{genero}', '{n_dueno}')")

                    cursor.execute(f'SELECT * FROM Pets WHERE nombre="{nombre}"')
                    get_nombreDb = cursor.fetchone()
                    conexion.commit()
                    conexion.close()
                    print(f'\n{self.white}[{self.cyan}!{self.white}] {self.lcyan}{nombre} {self.cyan}fue guardado en la BD correctamente.{self.reset}')
                    
                    opt = input(f'{self.white}[{self.magenta}>{self.white}] {self.reset}').lower()
                    if opt == 'x':
                        self.add_pet()
                except Exception as e:
                    print(e)

        else:
            self.start()

    def gen_report(self):
        """
        Generar reporte en Word con la lista de las mascotas
        """

        system('cls' if os.name == 'nt' else 'clear') # Si está en Windows usará 'cls', de lo contrario usará 'clear'
        system('title petSaver ~ Gen Report' if os.name == 'nt' else '') # Si está en Windows usará el CMD 'title'
        time.sleep(0.3)
        self.banner_genReport = f"""

                {self.white}[{self.cyan}Generar Reporte{self.white}]{self.lyellow}                                  {self.banner_cat_1}
                                                                  {self.banner_cat_2}
        {self.lyellow}Bienvenido, sigue las instrucciones                        {self.banner_cat_3}
        {self.lyellow}para generar un reporte en Word.                            {self.banner_cat_4}
                                                                    {self.banner_cat_5}
                                                                    {self.banner_cat_6}
                                                                    {self.banner_cat_7}
                                                                    {self.banner_cat_8}
                                                                    {self.banner_cat_9}
                    [x] Atrás                                       {self.banner_cat_10}
        """
        print(self.banner_genReport)
        name_docx = input(f'{self.white}[{self.lcyan}Nombre (sin el .docx){self.white}] {self.reset}')
        conn = sqlite3.connect(os.getenv('DB_NAME'))
        c = conn.cursor()

        table = 'Pets'

        c.execute(f'SELECT * FROM {table}')
        pets = c.fetchall()
        if pets: # Si existen mascotas en la base de datos, no hará nada
           pass
        else:
            print(f'{self.error_color}No hay ninguna mascota en la BD.')
            input("Presiona ENTER para volver.")
            self.main()
        

        doc = docx.Document()

        doc.add_heading('Reporte de Mascotas Registradas', level=1)

        for pet in pets:
            doc.add_paragraph('Nombre: '  + str(pet[1]))
            doc.add_paragraph('Edad: '  + str(pet[2]))
            doc.add_paragraph('Especie: '  + str(pet[3]))
            doc.add_paragraph('Género:  ' + str(pet[4]))
            doc.add_paragraph('Dueño: ' + str(pet[5]))
            doc.add_paragraph('╰───────╮╭───────╯')
        conn.close()
        doc.save(name_docx + '.docx')
        print(f'{self.white}{self.cyan}Reporte generado! {self.white}({self.lcyan}{name_docx}.docx{self.white}){self.reset}')
        option = input(f'\n{self.white}[{self.magenta}>{self.white}] {self.reset}').lower()
        if option == 'x':
                self.main()
        else:
            self.main()
    def main(self, run_start: bool=True):
        """
        Función principal de la aplicación.

        Si run_start es True, usará la funcion start()
        """
        
        system('cls' if os.name == 'nt' else 'clear') # Si está en Windows usará 'cls', de lo contrario usará 'clear'
        system('title petSaver ~ Main Menu' if os.name == 'nt' else '') # Si está en Windows usará el CMD 'title'
        if run_start:
            self.start()
        time.sleep(0.3)
        self.banner_main_ = f"""

                {self.white}[{self.lcyan}Main Menu{self.white}]{self.lcyan}                                            {self.banner_cat_1}
                                                                       {self.banner_cat_2}
            {self.white}[{self.lcyan}1{self.white}] {self.lcyan}Agregar Mascota                                        {self.banner_cat_3}
            {self.white}[{self.lcyan}2{self.white}] {self.lcyan}Generar Reporte                                        {self.banner_cat_4}
                                                                       {self.banner_cat_5}
            {self.white}[{self.lcyan}x{self.white}] {self.lcyan}Salir                                                  {self.banner_cat_6}
                                                                       {self.banner_cat_7}
                                                                       {self.banner_cat_8}
                                                                       {self.banner_cat_9}
                                                                       {self.banner_cat_10}

            
        """
        print(self.banner_main_)
        option = input(f'{self.white}[{self.magenta}>{self.white}] {self.reset}').lower()
        if option == '1':
            self.add_pet()
        elif option == '2':
            self.gen_report()
        elif option == 'x':
            sys.exit()
                
                        

if __name__ == '__main__':
    Pet = PetSaver()
    Pet.main()